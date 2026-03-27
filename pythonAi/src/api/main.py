# src/api/main.py
import sys
from pathlib import Path
root_dir = Path(__file__).parent.parent.parent
sys.path.insert(0, str(root_dir))

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import shutil
import tempfile
from contextlib import asynccontextmanager
from pydantic import BaseModel
from dotenv import load_dotenv
import os
import logging
import re
import uuid
import asyncio
from typing import Optional, AsyncGenerator
from fastapi import Request  
from fastapi.responses import StreamingResponse, FileResponse
from docx import Document
import json

from src.Agent.RAG_chain import RAGEngineLCEL
from src.Agent.agentCore import create_ai_agent, get_api_key, get_memory_as_langchain_messages
from src.Agent.memory import ConversationMemory
from src.Agent.DocumetLoader import DocumentProcessor  # 导入文档处理器

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局变量
rag_engine = None
agent = None
doc_processor = None  # 添加文档处理器
# 注意：生产环境应按会话管理memory，这里简单使用全局仅用于演示
memory = ConversationMemory()

# ---------- 数据模型 ----------
class ChatRequest(BaseModel):
    question: str
    session_id: str = "default"

class StreamRequest(BaseModel):
    question: str
    session_id: str = "default"
    temperature: float = 0.1
    use_knowledge: bool = True

class PolishRequest(BaseModel):
    text: str
    style: str = "professional"

# ---------- 回调处理器 ----------
from langchain_core.callbacks import BaseCallbackHandler

class DebugCallbackHandler(BaseCallbackHandler):
    def on_agent_action(self, action, **kwargs):
        logger.info(f"Agent Action: {action.log}")
    def on_agent_finish(self, finish, **kwargs):
        logger.info(f"Agent Finish: {finish.return_values}")

from langchain_core.runnables import RunnableConfig

# ---------- 生命周期管理 ----------
@asynccontextmanager
async def lifespan(app: FastAPI):
    global rag_engine, agent, doc_processor
    logger.info("正在初始化 RAG 引擎...")
    try:
        # 1. 初始化RAG引擎
        rag_engine = RAGEngineLCEL(
            persist_directory="./chroma_db",
            docs_path="./knowledge_base"
        )
        logger.info("✅ RAG 引擎初始化成功")
        
        # 2. 初始化文档处理器（传入rag_engine以便使用其embeddings）
        doc_processor = DocumentProcessor(rag_engine=rag_engine)
        logger.info("✅ 文档处理器初始化成功")
        
        # 3. 初始化Agent
        api_key = get_api_key()
        agent = create_ai_agent(api_key, rag_engine)
        
        # 4. 初始化工具（注入rag_engine和llm）
        # 注意：create_ai_agent内部会调用init_tools，所以这里不需要重复调用
        # 但如果create_ai_agent没有调用，需要在这里调用
        # init_tools(rag_engine, agent)  # 如果agent就是llm
        
        logger.info("✅ Agent 初始化成功")
    except Exception as e:
        logger.error(f"初始化失败: {e}")
        raise e
    yield
    logger.info("API 服务正在关闭")

app = FastAPI(
    title="简历润色助手 API",
    description="提供基于Agent的简历问答与润色服务",
    version="1.0.0",
    lifespan=lifespan
)

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- 文件处理辅助函数 ----------
async def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """
    从文件中提取文本内容（使用DocumentProcessor）
    """
    global doc_processor
    try:
        # 使用DocumentProcessor加载文件
        docs = doc_processor.load_file(file_path)
        # 合并所有文档的文本
        text = "\n".join([doc.page_content for doc in docs])
        return text
    except Exception as e:
        logger.error(f"从文件 {file_path} 提取文本失败: {e}")
        raise

#流式响应
async def stream_agent_response(
    question: str,
    file_text: str = "",
    temperature: float = 0.1
) -> AsyncGenerator[str, None]:
    """
    流式生成Agent响应的核心生成器 - 字符级别输出
    """
    global agent, memory
    
    try:
        # 1. 构造消息
        if file_text:
            agent_message = f"用户上传了文件，内容如下：\n{file_text}\n\n用户问题：{question}"
        else:
            agent_message = question
        
        # 2. 添加到记忆
        memory.add_to_memory('user', agent_message)
        langchain_messages = get_memory_as_langchain_messages(memory)
        
        # 3. 创建输入字典
        input_dict = {"messages": langchain_messages}
        
        # 4. 发送开始信号
        start_data = {"type": "start", "data": "开始处理..."}
        yield f"data: {json.dumps(start_data, ensure_ascii=False)}\n\n"
        
        full_response = ""
        buffer = ""  # 用于累积当前块
        
        # 使用 stream_mode="updates" 获取每一步的更新
        async for chunk in agent.astream(input_dict, stream_mode="updates"):
            for step_name, step_data in chunk.items():
                
                # 检查是否有消息
                if 'messages' in step_data and step_data['messages']:
                    last_msg = step_data['messages'][-1]
                    
                    # 处理内容块（LangGraph 格式）
                    if hasattr(last_msg, 'content_blocks'):
                        for block in last_msg.content_blocks:
                            if block['type'] == 'text':
                                text = block['text']
                                
                                # 检查是否是最终答案
                                has_tool_call = any(b['type'] == 'tool_call' for b in last_msg.content_blocks)
                                
                                if has_tool_call and text:
                                    # 工具调用前的思考 - 可以一次输出
                                    thinking_data = {"type": "thinking", "data": text}
                                    yield f"data: {json.dumps(thinking_data, ensure_ascii=False)}\n\n"
                                elif text:
                                    # 最终答案 - 逐字输出
                                    buffer += text
                                    # 当buffer累积到一定长度或遇到标点，可以输出
                                    # 但为了真正逐字，我们可以输出每个字符
                                    for char in text:
                                        chunk_data = {"type": "chunk", "data": char}
                                        yield f"data: {json.dumps(chunk_data, ensure_ascii=False)}\n\n"
                                        await asyncio.sleep(0.02)  # 控制输出速度
                                    full_response += text
                            
                            elif block['type'] == 'tool_call':
                                # 工具调用
                                tool_data = {
                                    "type": "tool", 
                                    "data": f"调用工具: {block['name']}"
                                }
                                yield f"data: {json.dumps(tool_data, ensure_ascii=False)}\n\n"
                    
                    # 处理普通的 AIMessageChunk
                    elif hasattr(last_msg, 'content') and last_msg.content:
                        text = last_msg.content
                        
                        if hasattr(last_msg, 'tool_calls') and last_msg.tool_calls:
                            # 工具调用
                            for tool_call in last_msg.tool_calls:
                                tool_data = {
                                    "type": "tool",
                                    "data": f"调用工具: {tool_call['name']}"
                                }
                                yield f"data: {json.dumps(tool_data, ensure_ascii=False)}\n\n"
                        else:
                            # 普通文本 - 逐字输出
                            for char in text:
                                chunk_data = {"type": "chunk", "data": char}
                                yield f"data: {json.dumps(chunk_data, ensure_ascii=False)}\n\n"
                                await asyncio.sleep(0.02)  # 控制输出速度
                            full_response += text
        
        # 5. 发送完成信号
        done_data = {"type": "done", "data": full_response}
        yield f"data: {json.dumps(done_data, ensure_ascii=False)}\n\n"
        
        # 6. 将完整回复存入记忆
        memory.add_to_memory('assistant', full_response)
        
    except Exception as e:
        logger.error(f"流式生成错误: {e}", exc_info=True)
        error_data = {"type": "error", "data": f"生成回答时出错: {str(e)}"}
        yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"
# ---------- API端点 ----------
@app.post("/ask")
async def ask_endpoint(
    raw_request: Request,
    file: UploadFile = File(None),
    question: str = Form(None)
):
    """
    统一的问答接口（非流式）
    """
    logger.info(f"=== /ask 请求 ===")
    
    global agent, memory
    if agent is None or rag_engine is None:
        raise HTTPException(status_code=503, detail="服务未就绪")

    user_input = ""
    file_text = ""
    has_file = file is not None and file.filename is not None

    if has_file:
        # 文件上传处理
        if not question:
            raise HTTPException(status_code=400, detail="请提供问题")
        user_input = question.strip()
        
        # 验证文件类型
        allowed_extensions = ['.txt', '.pdf', '.docx', '.md']
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_ext}")

        # 保存临时文件并提取文本
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                shutil.copyfileobj(file.file, tmp)
                tmp_path = tmp.name
            file_text = await extract_text_from_file(tmp_path, file_ext)
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            await file.close()
    else:
        # 纯文本JSON处理
        try:
            body = await raw_request.json()
            user_input = body.get("question", "").strip()
            if not user_input:
                raise HTTPException(status_code=400, detail="字段 'question' 不能为空")
        except Exception as e:
            logger.error(f"JSON 解析失败: {e}")
            raise HTTPException(status_code=400, detail="无效的 JSON 格式")

    # 构造Agent消息
    if file_text:
        agent_message = f"用户上传了文件，内容如下：\n{file_text}\n\n用户问题：{user_input}"
    else:
        agent_message = user_input

    # 调用Agent
    memory.add_to_memory('user', agent_message)
    langchain_messages = get_memory_as_langchain_messages(memory)

    try:
        config = RunnableConfig(callbacks=[DebugCallbackHandler()])
        input_dict = {"messages": langchain_messages}
        
        # 同步调用（非流式）
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, lambda: agent.invoke(input_dict, config=config))
        
        # 提取答案
        if hasattr(response, 'content'):
            answer = response.content
        elif isinstance(response, dict) and 'messages' in response:
            last_message = response['messages'][-1]
            answer = last_message.content if hasattr(last_message, 'content') else str(last_message)
        else:
            answer = str(response)
        
        memory.add_to_memory('assistant', answer)
        return {"question": user_input, "answer": answer}
        
    except Exception as e:
        logger.error(f"Agent调用失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="处理失败")

@app.post("/ask/stream")
async def ask_stream_endpoint(
    request: Request,
    file: UploadFile = File(None)
):
    """
    完整的流式问答接口
    """
    logger.info("=== /ask/stream 请求 ===")
    
    # 1. 解析请求
    try:
        body = await request.json()
        question = body.get("question", "")
        session_id = body.get("session_id", "default")
        temperature = body.get("temperature", 0.1)
        use_knowledge = body.get("use_knowledge", True)
    except Exception as e:
        logger.error(f"JSON 解析失败: {e}")
        raise HTTPException(status_code=400, detail=f"无效的JSON: {e}")
    
    if not question:
        raise HTTPException(status_code=400, detail="缺少question字段")
    
    # 2. 检查服务状态
    if agent is None or rag_engine is None:
        raise HTTPException(status_code=503, detail="服务未就绪")
    
    # 3. 处理文件上传（如果有）
    file_text = ""
    if file and file.filename:
        try:
            allowed_extensions = ['.txt', '.pdf', '.docx', '.md']
            file_ext = Path(file.filename).suffix.lower()
            if file_ext not in allowed_extensions:
                raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_ext}")
            
            tmp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                    shutil.copyfileobj(file.file, tmp)
                    tmp_path = tmp.name
                file_text = await extract_text_from_file(tmp_path, file_ext)
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                await file.close()
        except Exception as e:
            logger.error(f"文件处理失败: {e}")
            raise HTTPException(status_code=400, detail=f"文件处理失败: {str(e)}")
    
    # 4. 返回流式响应
    return StreamingResponse(
        stream_agent_response(
            question=question,
            file_text=file_text,
            temperature=temperature
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

@app.post("/download_docx", response_class=FileResponse)
async def download_docx(request: PolishRequest, background_tasks: BackgroundTasks):
    """生成Word文档并返回"""
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="文本不能为空")

    doc = Document()
    doc.add_heading('润色后的文本', level=1)
    for para in request.text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    file_name = f"润色结果_{uuid.uuid4().hex[:8]}.docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    background_tasks.add_task(os.unlink, tmp_path)

    return FileResponse(
        path=tmp_path,
        filename=file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/health")
async def health_check():
    """健康检查"""
    return {
        "status": "healthy",
        "rag_engine": rag_engine is not None,
        "agent": agent is not None,
        "doc_processor": doc_processor is not None
    }